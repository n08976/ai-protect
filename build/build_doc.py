"""Build the companion technical doc to v2.1.
Audience: offensive security leads (the five vertical owners).
Mirrors v2.1's styling: 1x1 shaded callout tables (light-blue takeaway, orange callout),
navy-banded data tables, Heading 1/2/3 hierarchy, embedded PNG diagrams.
"""

import os

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DIAGRAMS = os.path.join(REPO, "diagrams")
DOCS = os.path.join(REPO, "docs")
os.makedirs(DOCS, exist_ok=True)

# Palette mirrors v2.1
NAVY = "1F3A5F"
NAVY_RGB = RGBColor(0x1F, 0x3A, 0x5F)
WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)
TAKEAWAY = "EAF3FA"
CALLOUT = "FFF4E5"
ALT_ROW = "F2F5F9"
ACCENT = "C04A2B"
ACCENT_RGB = RGBColor(0xC0, 0x4A, 0x2B)

# ---------- helpers ----------

def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_cell_borders(cell, color="BFC8D6", size="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def remove_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_para(doc, text, *, bold=False, italic=False, size=None, color=None, align=None,
             style="Normal", space_before=None, space_after=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if size is not None:
        r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = NAVY_RGB
    return h


def add_bullets(doc, items, bold_lead=False):
    for it in items:
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.left_indent = Inches(0.25)
        # bullet formatting
        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numPr.append(ilvl)
        numId = OxmlElement("w:numId")
        numId.set(qn("w:val"), "1")
        numPr.append(numId)
        pPr.append(numPr)
        # text
        if bold_lead and " — " in it:
            head, tail = it.split(" — ", 1)
            r1 = p.add_run(head + " — ")
            r1.bold = True
            p.add_run(tail)
        else:
            p.add_run(it)
    return


def add_callout(doc, kind, title, body_lines):
    """kind: 'takeaway' (light blue) or 'callout' (orange).
    body_lines: list of bullet strings, or single string."""
    color = TAKEAWAY if kind == "takeaway" else CALLOUT
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    cell = t.rows[0].cells[0]
    set_cell_shading(cell, color)
    set_cell_borders(cell, color="BFC8D6", size="4")
    # set width
    cell.width = Inches(6.5)
    tcW = cell._tc.get_or_add_tcPr()
    tcW2 = OxmlElement("w:tcW")
    tcW2.set(qn("w:w"), "9360")  # 6.5 inches in twips
    tcW2.set(qn("w:type"), "dxa")
    tcW.append(tcW2)
    # margins
    for tag, val in (("top", "120"), ("left", "180"), ("bottom", "120"), ("right", "180")):
        m = OxmlElement(f"w:{tag}")
        m.set(qn("w:w"), val)
        m.set(qn("w:type"), "dxa")
        tcMar = cell._tc.tcPr.find(qn("w:tcMar")) or OxmlElement("w:tcMar")
        tcMar.append(m)
        if tcMar.getparent() is None:
            cell._tc.tcPr.append(tcMar)
    # title paragraph
    cell.text = ""  # clear
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = NAVY_RGB
    if isinstance(body_lines, str):
        body_lines = [body_lines]
    for ln in body_lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        # bullet character ourselves rather than list style (cleaner inside cell)
        if len(body_lines) > 1:
            r = p.add_run("•  " + ln)
        else:
            r = p.add_run(ln)
        r.font.size = Pt(10.5)
    add_para(doc, "", space_after=2)
    return t


def _set_fixed_table_layout(t, col_widths_in, tight_cells=False):
    """Force fixed layout with explicit tblGrid and tblW so cell widths actually apply."""
    tbl = t._tbl
    tblPr = tbl.tblPr
    # tblLayout fixed
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    # tighten default cell margins on the table
    if tight_cells:
        for old in tblPr.findall(qn("w:tblCellMar")):
            tblPr.remove(old)
        tcm = OxmlElement("w:tblCellMar")
        for tag, val in (("top", "30"), ("bottom", "30"), ("left", "60"), ("right", "60")):
            m = OxmlElement(f"w:{tag}")
            m.set(qn("w:w"), val)
            m.set(qn("w:type"), "dxa")
            tcm.append(m)
        tblPr.append(tcm)
    # tblW = sum
    total_twips = int(sum(col_widths_in) * 1440)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(total_twips))
    tblW.set(qn("w:type"), "dxa")
    # remove existing tblW if any
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblPr.append(tblW)
    # tblGrid
    for existing in tbl.findall(qn("w:tblGrid")):
        tbl.remove(existing)
    grid = OxmlElement("w:tblGrid")
    for w in col_widths_in:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w * 1440)))
        grid.append(gc)
    # tblGrid must come right after tblPr
    tblPr.addnext(grid)
    # also set each cell width
    for row in t.rows:
        for ci, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(int(col_widths_in[ci] * 1440)))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


def _set_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    cs = OxmlElement("w:cantSplit")
    trPr.append(cs)


def add_data_table(doc, headers, rows, col_widths=None, header_color=NAVY, alt_color=ALT_ROW,
                   small=False):
    nrows = len(rows) + 1
    ncols = len(headers)
    t = doc.add_table(rows=nrows, cols=ncols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    # apply fixed layout if widths supplied
    if col_widths:
        _set_fixed_table_layout(t, col_widths, tight_cells=small)
    body_size = Pt(9) if small else Pt(10)
    head_size = Pt(10) if small else Pt(10.5)
    # headers
    for ci, h in enumerate(headers):
        c = t.rows[0].cells[ci]
        set_cell_shading(c, header_color)
        set_cell_borders(c, color="FFFFFF", size="4")
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = head_size
        r.font.color.rgb = WHITE_RGB
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_row_cant_split(t.rows[0])
    # rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            shade = alt_color if ri % 2 == 0 else "FFFFFF"
            set_cell_shading(c, shade)
            set_cell_borders(c, color="DDE3EB", size="4")
            c.text = ""
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else (WD_ALIGN_PARAGRAPH.CENTER if small else WD_ALIGN_PARAGRAPH.LEFT)
            r = p.add_run(str(val))
            r.font.size = body_size
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_row_cant_split(t.rows[ri + 1])
    add_para(doc, "", space_after=2)
    return t


def add_image(doc, path, width_in=6.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run()
    r.add_picture(path, width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(10)
        cr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def configure_doc(doc):
    # default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    # heading colors
    for lvl in range(1, 4):
        try:
            h = doc.styles[f"Heading {lvl}"]
            h.font.color.rgb = NAVY_RGB
            h.font.name = "Calibri"
        except KeyError:
            pass
    # margins
    for sec in doc.sections:
        sec.top_margin = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin = Inches(0.9)
        sec.right_margin = Inches(0.9)


# ============================================================
# BUILD
# ============================================================

doc = Document()
configure_doc(doc)

# ----- Cover page -----
cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_title.paragraph_format.space_before = Pt(120)
ct = cover_title.add_run("AI Security Pipeline")
ct.bold = True
ct.font.size = Pt(34)
ct.font.color.rgb = NAVY_RGB

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(40)
sr = sub.add_run("Technical Companion to the v2.1 Operating Model")
sr.font.size = Pt(18)
sr.font.color.rgb = NAVY_RGB

# Decorative bar
bar = doc.add_paragraph()
bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
br = bar.add_run("─" * 30)
br.font.color.rgb = ACCENT_RGB
br.font.size = Pt(14)

audience = doc.add_paragraph()
audience.alignment = WD_ALIGN_PARAGRAPH.CENTER
audience.paragraph_format.space_before = Pt(60)
ar = audience.add_run("Audience · Offensive Security Leadership Team")
ar.font.size = Pt(13)
ar.bold = True
ar.font.color.rgb = NAVY_RGB

audience2 = doc.add_paragraph()
audience2.alignment = WD_ALIGN_PARAGRAPH.CENTER
ar2 = audience2.add_run("Application Security  ·  Threat Intelligence  ·  Threat Hunt  ·  Red Team  ·  Security Control Validation")
ar2.font.size = Pt(11)
ar2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver.paragraph_format.space_before = Pt(60)
vr = ver.add_run("Version 1.0 · Companion to v2.1 · Internal — Offensive Security")
vr.font.size = Pt(10)
vr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

prep = doc.add_paragraph()
prep.alignment = WD_ALIGN_PARAGRAPH.CENTER
pr_ = prep.add_run("Prepared by the Office of the Director, Offensive Security")
pr_.font.size = Pt(10)
pr_.italic = True
pr_.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

page_break(doc)

# ----- Executive summary for the leads -----
add_heading(doc, "Executive Summary for the Leads", level=1)

add_para(doc,
    "The v2.1 operating model committed our organization to a shift from gatekeeper to embedded "
    "advisor, anchored on a paved-road infrastructure (Claude as primary sanctioned model, an MCP "
    "server farm, a scoped agent runtime), risk-tiered engagement, and a phased capability build. "
    "Executive sponsorship for that direction is in hand. This companion document is the engineering "
    "blueprint your teams will operate against."
)

add_para(doc,
    "It does three things. First, it lays out the eight-stage AI assurance pipeline that "
    "operationalizes v2.1 — what runs at each stage, which tools we use, and what the inputs and "
    "outputs are. Second, it translates each stage into per-vertical capability builds for the five "
    "functions you each lead, phased to honor v2.1's modesty constraint in Phase 1. Third, it "
    "extends the v2.1 RACI with eighteen pipeline-specific activities so that ownership is "
    "unambiguous from day one."
)

add_callout(doc, "takeaway", "Key takeaways", [
    "This is the technical layer beneath v2.1, not a competing proposal — every architectural choice traces back to commitments already approved.",
    "Eight-stage pipeline anchored on the v2.1 sanctioned infrastructure: Claude gateway, MCP farm, agent runtime.",
    "Tooling for Phase 1 is deliberately conservative: open-source PoCs (garak, PyRIT, Semgrep, Trivy, ModelScan) operated by existing headcount.",
    "Each of the five verticals owns clear primary territory in the pipeline; the ownership matrix is in Section 5.",
    "Two dashboard surfaces — technical for operators, executive for the CISO — feed from a single findings store with HIPAA/HITRUST evidence pre-mapped.",
    "Eighteen-row RACI extension slots cleanly into v2.1's nine owner columns; no realignment required.",
])

add_para(doc, "")
add_image(doc, os.path.join(DIAGRAMS, "02_v21_mapping.png"), width_in=6.5,
          caption="Figure 1. Where this companion sits relative to the v2.1 operating model.")

page_break(doc)

# ----- Section 1: How this fits with v2.1 -----
add_heading(doc, "1. How This Companion Fits with v2.1", level=1)

add_para(doc,
    "v2.1 was written for the CISO and cyber executives. It described what the offensive security "
    "team will do under AI transformation and why — the strategic reframe to empirical truth "
    "function, the risk tiering framework, the engagement model, the phased roadmap, and the RACI "
    "across nine owner columns. It described the architectural commitments at the level of "
    "principle: foundation model usage funneled through a sanctioned gateway, MCP servers in a "
    "curated registry, agents in a scoped runtime, telemetry unified."
)

add_para(doc,
    "This companion is for the offensive security leads. It describes how those commitments get "
    "engineered into a working pipeline — the eight stages, the tools at each stage, who runs them, "
    "what is logged, how findings are normalized, how dashboards are populated, and what each of "
    "your teams builds in each of the three phases."
)

add_heading(doc, "Inheritance Map", level=2)

add_data_table(doc,
    headers=["v2.1 commitment", "How this document operationalizes it"],
    rows=[
        ["Paved road over gates",
         "Sanctioned LLM gateway and curated MCP registry are the only approved foundation-model and tool paths. Stages 0–3 of the pipeline test conformance; deviation is a finding."],
        ["Risk-tiered engagement",
         "Stage 1 is automated tier scoring on intake. Tiers determine which downstream stages are mandatory, optional, or async. Tier 1 has hard pre-deploy gates; Tier 4 is self-service."],
        ["Continuous validation",
         "Stages 6 and 7 implement re-scan cadence, drift detection, and replay of new jailbreaks/CVEs against the live portfolio."],
        ["Five-function realignment",
         "Section 5 spells out per-vertical builds for AppSec, TI, TH, RT, SCV — including which pipeline stages each owns."],
        ["RACI across nine owners",
         "Section 8 adds eighteen pipeline activities to the existing RACI columns. No new column owners introduced."],
        ["Phase 1 modesty",
         "Phase 1 tooling is open-source or already-licensed. No headcount asks. Demo red team is the artifact that earns Phase 2."],
        ["Sanctioned infrastructure",
         "Pipeline stages instrument the gateway, MCP farm, agent runtime, and data plane already named in v2.1. No parallel stack."],
        ["Empirical truth function",
         "Pipeline output feeds the quarterly board/risk-committee report; findings are queryable as HIPAA/HITRUST evidence."],
    ],
    col_widths=[2.2, 4.2])

page_break(doc)

# ----- Section 2: Pipeline architecture overview -----
add_heading(doc, "2. Pipeline Architecture Overview", level=1)

add_para(doc,
    "The pipeline is eight stages, each with a defined input, output, primary owner, and tool set. "
    "Stages run as workflows on a shared orchestration layer; findings are normalized to OCSF and "
    "stored in a single findings system; downstream notification, remediation, and dashboard surfaces "
    "are the consumers."
)

add_image(doc, os.path.join(DIAGRAMS, "01_pipeline_overview.png"), width_in=6.6,
          caption="Figure 2. The eight-stage pipeline, anchored on the v2.1 sanctioned infrastructure.")

add_heading(doc, "Three Layers, One Pipeline", level=2)

add_bullets(doc, [
    "Stage layer — eight discrete assurance stages from discovery to reporting. Each stage is a workflow with its own SLA and pass/fail criteria.",
    "Orchestration & data plane — Argo Workflows (or Tekton) for execution, Kafka for events, DefectDojo for findings, Vault for secrets, OPA for deploy gates. This layer is operated by Platform Engineering; offensive security defines policy.",
    "Sanctioned infrastructure — the v2.1 gateway, MCP farm, agent runtime, data plane, and telemetry mesh are not part of the pipeline; they are what the pipeline tests, monitors, and reports on.",
], bold_lead=True)

add_callout(doc, "callout", "The architectural rule that keeps this honest",
    "If an AI asset does not flow through the gateway and use a registered MCP, the pipeline cannot see it. "
    "Off-pipeline deployment is therefore treated as an incident, not a finding. Stage 0 (discovery) is the "
    "control that makes the rest of the pipeline credible."
)

add_heading(doc, "What Travels on the Wire", level=2)

add_para(doc,
    "Findings produced by any stage land in a single store with a uniform schema. This is the design "
    "decision that makes everything downstream tractable: dashboards, RACI accountability, SLA tracking, "
    "compliance evidence, and threat-hunt enrichment all consume the same normalized records."
)

add_data_table(doc,
    headers=["Field", "Source", "Why it matters"],
    rows=[
        ["AIID", "Stage 0 — CMDB tag", "Single key to join findings, telemetry, and ticketing per asset."],
        ["Tier", "Stage 1 — automated scoring", "Drives SLA, mandatory stages, and dashboard segmentation."],
        ["Severity (CVSS or AI-adjusted)", "Each scanner", "Normalized to a 0–10 scale; AI-specific adjustments documented in Stage 4."],
        ["ATLAS technique", "Stage 4 — red team output", "Coverage tracking and adversary-emulation reporting."],
        ["HIPAA/HITRUST control", "Auto-mapped at finding creation", "Audit evidence is a query, not a fire drill."],
        ["BU / owner email", "Stage 0 intake + IdP lookup", "Notification routing and SLA accountability."],
        ["Telemetry references", "Stage 6", "Link runtime evidence (prompt/tool-call logs) to the finding."],
    ],
    col_widths=[1.8, 1.7, 3.0])

page_break(doc)

# ----- Section 3: Stage-by-stage -----
add_heading(doc, "3. Stage-by-Stage Detail", level=1)

stages_detail = [
    {
        "n": "0",
        "name": "Discovery & Intake",
        "purpose": "Establish the denominator. Every AI asset gets registered with an AIID, owner, and risk inputs before any other stage runs. Shadow-AI is detected and routed into intake.",
        "inputs": "Self-service portal submissions; CASB egress telemetry; GitHub/GitLab org scans; Microsoft Purview / Copilot Studio inventory; browser-extension paste-PHI signals.",
        "outputs": "AIID record in CMDB; intake form payload; tier inputs ready for Stage 1.",
        "tools": ["ServiceNow (intake form + record)", "CASB / SWG egress monitoring", "GitHub org scanner (Semgrep + custom signatures for `openai`, `anthropic`, `langchain`, `llama-index`)", "Purview / Copilot Studio API", "Browser-extension telemetry pipeline"],
        "primary": "Threat Intel (shadow-AI discovery); AppSec (intake portal owner)",
        "support": "All five verticals consume the inventory",
        "sla": "Intake → AIID issued within 1 business day for active submissions; passive discovery within 24h of egress signal",
        "hc": "BAA-coverage flag is set at intake. If a non-BAA'd third-party AI service is detected in egress, this is a same-day escalation, not a quarterly report item.",
        "raci": "Activity 'Discovery & intake of AI assets' (Section 8) — A: Threat Intel; R: TI + AppSec; C: PlatEng, Privacy; I: AI Gov, Biz Owner",
    },
    {
        "n": "1",
        "name": "Triage & Tiering",
        "purpose": "Convert intake inputs into a deterministic tier (1–4) and a downstream engagement plan. No human judgment at this stage — humans appeal, they don't assign.",
        "inputs": "Intake form (data classes, decision impact, integrations, user population); discovery signals (egress patterns); BAA flag.",
        "outputs": "Tier assigned; mandatory stage list set; SLA clock started; ticket auto-opened in vertical's queue.",
        "tools": ["OPA-based tiering policy (versioned, audited)", "CMDB tag update", "Jira/ServiceNow auto-ticket creation", "Stage-routing in orchestrator"],
        "primary": "AppSec (owns the tiering policy implementation, AI Governance owns the framework)",
        "support": "Privacy/Compliance review of policy quarterly",
        "sla": "Auto-tier within 15 minutes of intake completion",
        "hc": "Any asset that touches PHI, ePHI, claims, clinical decisioning, or has external-facing surface is Tier 1 by rule. Tier reductions for Tier 1 candidates require AI Governance + Privacy sign-off, recorded in CMDB.",
        "raci": "Activity 'Tier assignment for new AI asset' — A: AI Gov; R: AppSec; C: Privacy/Compliance; I: TI, RT, Biz Owner",
    },
    {
        "n": "2",
        "name": "Static / Pre-Production Scanning",
        "purpose": "Catch what's catchable without running the system. Code, dependencies, containers, IaC, model artifacts, prompt templates, and RAG corpus all get scanned before any dynamic stage runs.",
        "inputs": "Source repository, build artifacts, model artifacts (HuggingFace or internal), RAG corpus URIs, IaC.",
        "outputs": "Findings with severity, CWE/CVE/ATLAS mapping, file/line references, suggested fix where available.",
        "tools": ["Semgrep + CodeQL (custom rules for unsafe prompt interpolation, agent tool-call sinks)", "TruffleHog / gitleaks", "Trivy / Grype (containers + SBOM)", "Checkov / tfsec / KICS (IaC)", "ModelScan (Protect AI) for serialized-model RCE", "Sigstore signature verification on model artifacts", "Presidio + clinical NER on RAG corpus for PHI", "Custom corpus poisoning detector (hash baseline + drift)"],
        "primary": "AppSec",
        "support": "SCV validates that scanners are in fact running on every asset of the relevant tier",
        "sla": "Tier 1: blocking; Tier 2: blocking with waiver path; Tier 3–4: async, ticketed",
        "hc": "ModelScan is non-negotiable on any HuggingFace import path. Pickle-based model RCE is one of the highest-severity supply-chain risks in the AI stack and remains under-addressed in most healthcare orgs.",
        "raci": "Activity 'Static and supply-chain scanning' — A: AppSec; R: AppSec; C: PlatEng (CI/CD integration); I: SCV, Biz Owner",
    },
    {
        "n": "3",
        "name": "Dynamic AppSec",
        "purpose": "Test the running system from outside. Web/API surfaces, authenticated flows, fuzz of OpenAPI contracts, exposed CVEs.",
        "inputs": "Deployed asset URL, auth credentials (rotated test creds), OpenAPI spec where available.",
        "outputs": "Findings with reproduction steps and severity.",
        "tools": ["Burp Suite Enterprise (API-driven, stored session macros)", "Nuclei (CVE / exposure sweeps)", "OWASP ZAP (parallel/redundant for OSS coverage)", "Schemathesis / RESTler for OpenAPI fuzz", "Nessus / Tenable on supporting infrastructure"],
        "primary": "AppSec",
        "support": "Red Team contributes campaign-mode dynamic testing for Tier 1; SCV validates coverage",
        "sla": "Tier 1: pre-deploy + weekly; Tier 2: pre-deploy + monthly; Tier 3–4: monthly automated",
        "hc": "Authenticated scans against assets that touch PHI must use synthetic test data only. The pipeline enforces a separate test-tenant isolation boundary; any scanner that leaks production PHI into a finding becomes the next incident.",
        "raci": "Activity 'Dynamic application & API testing' — A: AppSec; R: AppSec, RT (campaigns); C: PlatEng; I: SCV, Biz Owner",
    },
    {
        "n": "4",
        "name": "AI Red Team",
        "purpose": "Test the AI system as an adversarial system. Jailbreaks, prompt injection (direct and indirect), tool/agent abuse, adversarial ML, and healthcare-specific harm probes. This is the differentiator stage and is detailed in Section 4.",
        "inputs": "Deployed asset access via gateway; scoped test identity; permission to inject content into RAG corpus (test partition only).",
        "outputs": "Findings with attack chains, ATLAS technique IDs, regression-test cases (PromptFoo) for any successful exploit, and a campaign report for Tier 1.",
        "tools": ["garak (NVIDIA) — baseline + custom probes", "PyRIT (Microsoft) — multi-turn orchestrators", "ART / Counterfit / TextAttack — adversarial ML", "PromptFoo — regression suite", "Custom indirect-injection corpus", "Tool-call fuzz harness (in-house)"],
        "primary": "Red Team",
        "support": "AppSec (handles findings into remediation); Threat Intel (provides current jailbreak techniques); SCV (validates that successful exploits are then blocked by guardrails)",
        "sla": "Tier 1: pre-deploy + monthly; Tier 2: pre-deploy + quarterly; Tier 3: annual or on material change",
        "hc": "Healthcare custom probe set runs every campaign: PHI elicitation, clinical hallucination, off-label drug recs, minimum-necessary violation, BAA-bypass tool calls. These are not optional add-ons — they are the probes that justify our existence.",
        "raci": "Activity 'AI red team campaign' — A: RT; R: RT; C: AppSec, TI, SCV, Privacy; I: AI Gov, Biz Owner",
    },
    {
        "n": "5",
        "name": "Remediation",
        "purpose": "Close findings. Auto where possible, ticketed where not, with WAF/guardrail backstops where root-cause fix takes longer than the SLA.",
        "inputs": "Findings from any prior stage.",
        "outputs": "Closed findings (with evidence); active backstops (WAF rules, guardrail policies); waivers (time-bound, accountable).",
        "tools": ["Auto-PR system for dependency bumps and IaC fixes", "Llama Guard 3 + NeMo Guardrails (guardrail injection at gateway)", "LLM Guard / Rebuff (input-side prompt-injection defense)", "Presidio (output PHI redaction)", "Kong / Apigee WAF rule push", "DefectDojo waiver workflow with hard expiry"],
        "primary": "AppSec (root-cause); SCV (validates closure)",
        "support": "RT regression-tests successful exploits in PromptFoo; Threat Hunt monitors for re-emergence",
        "sla": "Tier 1 critical: 7 days; high: 14 days; medium: 30 days; low: 90 days. Backstops within 24h for Tier 1 critical.",
        "hc": "Output-side PHI redaction (Presidio + clinical NER) is layered — never the only control. A model that hallucinates PHI is still a privacy risk even if the output filter catches it; the root cause must be addressed.",
        "raci": "Activity 'Remediation & closure validation' — A: AppSec; R: AppSec, SCV (validation); C: RT; I: Biz Owner, AI Gov",
    },
    {
        "n": "6",
        "name": "Continuous Monitoring",
        "purpose": "Production-time assurance. Runtime telemetry, drift detection, scheduled re-scan, replay of new attacks against the live portfolio.",
        "inputs": "Telemetry mesh (prompts, completions, tool calls, retrieval queries) — PHI-redacted at the edge; new threat-intel signals.",
        "outputs": "Hunt-ready data; drift alerts; periodic re-scan results; new findings from intel-driven replay.",
        "tools": ["OpenSearch / SIEM ingest", "Custom drift detectors on jailbreak success rate, refusal rate, tool-call distribution", "Cron'd garak/PyRIT runs (weekly Tier 1, monthly Tier 2)", "Threat-feed → asset replay job", "Anomaly detection on prompt/tool-call sequences"],
        "primary": "Threat Hunt (detection hypotheses + analytics); SCV (re-scan cadence)",
        "support": "Threat Intel (feeds new techniques); Red Team (validates that new techniques don't exploit our portfolio)",
        "sla": "Telemetry latency to hunt-ready: <15 minutes. Re-scan cadence per tier. Threat-driven replay within 48h of credible new-jailbreak intel.",
        "hc": "Telemetry is PHI-redacted at the edge before storage. The pipeline enforces this — un-redacted telemetry never reaches the long-term store. SIEM ingest is the secondary control, not the primary.",
        "raci": "Activity 'Continuous monitoring & drift detection' — A: TH; R: TH, SCV; C: TI, RT, PlatEng; I: AppSec, Biz Owner",
    },
    {
        "n": "7",
        "name": "Reporting & Notification",
        "purpose": "Close the loop with the asset owner and produce the evidence the team needs for the seat at the table — operator dashboards, executive dashboards, compliance evidence, and the quarterly board report.",
        "inputs": "Normalized findings, telemetry summaries, RACI ownership, SLA clocks.",
        "outputs": "Per-asset Security Report Card (PDF + portal); Slack/Teams DM + email; Jira/ServiceNow ticket; dashboard refresh; board-ready quarterly export.",
        "tools": ["DefectDojo report generation", "Slack/Teams notifier", "Jira/ServiceNow integration", "Grafana (technical dashboard)", "Superset / Power BI (executive dashboard)", "HIPAA/HITRUST evidence query templates"],
        "primary": "AppSec (asset-level reports); SCV (control reports); Office of the Director (board export)",
        "support": "All verticals contribute their stage-level data",
        "sla": "Notification within 15 minutes of finding creation. Dashboard freshness < 5 minutes. Quarterly board report within 5 business days of quarter close.",
        "hc": "The board report is the artifact that turns the seat at the table from a claim into a habit. It is the single most important reporting deliverable from this pipeline. Owner attestation that controls were validated, not just that scanners ran, is what differentiates this from a vulnerability-count report.",
        "raci": "Activity 'Asset Security Report Card' — A: AppSec; R: AppSec, SCV; C: All other verticals; I: Biz Owner. Activity 'Quarterly board / risk-committee report' — A: Director; R: All vertical leads; C: AI Gov, Privacy; I: CISO.",
    },
]

for st in stages_detail:
    add_heading(doc, f"Stage {st['n']} — {st['name']}", level=2)

    add_para(doc, "Purpose. " + st["purpose"])

    add_heading(doc, "Inputs / Outputs", level=3)
    add_bullets(doc, [
        f"Inputs — {st['inputs']}",
        f"Outputs — {st['outputs']}",
    ], bold_lead=True)

    add_heading(doc, "Tooling", level=3)
    add_bullets(doc, st["tools"])

    add_heading(doc, "Ownership and SLA", level=3)
    add_bullets(doc, [
        f"Primary owner — {st['primary']}",
        f"Support — {st['support']}",
        f"SLA / cycle time — {st['sla']}",
        f"v2.1 RACI — {st['raci']}",
    ], bold_lead=True)

    add_callout(doc, "callout", "Healthcare consideration", st["hc"])
    page_break(doc)

# ----- Section 4: AI Red Team capability deep-dive -----
add_heading(doc, "4. AI Red Team Capability Deep-Dive", level=1)

add_para(doc,
    "Stage 4 is the differentiator. Most healthcare organizations can stand up Stages 0–3 and Stage "
    "6 with existing application security tooling and skill. Stage 4 is what proves we are doing AI "
    "security and not appsec-with-an-AI-decal — and it is where the empirical-truth claim from v2.1 "
    "is earned or lost."
)

add_image(doc, os.path.join(DIAGRAMS, "03_ai_redteam_killchain.png"), width_in=6.6,
          caption="Figure 3. The AI red-team kill chain — probes, tooling, and ATLAS coverage at each step.")

add_heading(doc, "What We Run, in Order", level=2)

add_para(doc,
    "Each Tier 1 campaign walks the kill chain end-to-end. Lower tiers may stop after Stage 4.b "
    "(single-turn) for cycle-time reasons, with the remainder run on quarterly or annual cadence."
)

add_heading(doc, "4.a  Reconnaissance", level=3)
add_bullets(doc, [
    "Model fingerprinting via response-signature probes (garak `lmrc` family).",
    "System-prompt leak attempts (`leakreplay`).",
    "Endpoint and tool-call surface enumeration via gateway logs (read-only).",
    "Output: a target profile that informs the rest of the campaign.",
])

add_heading(doc, "4.b  Single-turn jailbreak", level=3)
add_bullets(doc, [
    "garak baseline probes: `promptinject`, `dan`, `encoding`, `leakreplay`, `realtoxicityprompts`, `xss`, `malwaregen`.",
    "Custom probes for healthcare context — see callout below.",
    "PromptFoo regression suite: every successful exploit becomes a permanent test case.",
])

add_heading(doc, "4.c  Multi-turn attack", level=3)
add_bullets(doc, [
    "PyRIT orchestrators: single → multi-turn → crescendo → red-team-bot, escalating until refusal or budget.",
    "Custom converters for medical-context jailbreaks (clinician role-play, patient-history justifications).",
    "Drift testing: does the model become more compliant over a long session?",
])

add_heading(doc, "4.d  Indirect prompt injection", level=3)
add_bullets(doc, [
    "Curated poisoned-document corpus seeded into a test partition of the asset's RAG store.",
    "Email/document/PDF payloads for assets that ingest external content.",
    "Tool-result hijack: payload returned by a downstream tool that the agent then acts on.",
    "These are the attacks most likely to compromise a Tier 1 agent in production. Coverage here is non-negotiable.",
])

add_heading(doc, "4.e  Tool / agent abuse", level=3)
add_bullets(doc, [
    "Tool-call fuzzing — does the agent call tools it shouldn't, or call them with arguments outside the schema?",
    "Confused-deputy: get the agent to use its privileges on attacker's behalf.",
    "Excessive agency (OWASP LLM06) — chain of tool calls that escalates beyond the intended task.",
    "Tool-scoping at registration is the highest-leverage control here. This stage validates the scoping decision.",
])

add_heading(doc, "4.f  Adversarial ML", level=3)
add_bullets(doc, [
    "IBM ART, Counterfit, TextAttack against any classifier or embedding model in the loop.",
    "Embedding poisoning / inversion attacks on the vector store.",
    "Model extraction probes against any self-hosted model.",
])

add_heading(doc, "4.g  Persistence / exfiltration", level=3)
add_bullets(doc, [
    "Memory poisoning of agent context store.",
    "PHI extraction from training/fine-tuning data — relevant for any internally fine-tuned model.",
    "Off-pipeline pivot: can the agent be coerced into using a non-sanctioned model or tool?",
])

add_callout(doc, "callout", "Healthcare custom probe set — runs on every Tier 1 campaign", [
    "PHI elicitation — direct and indirect attempts to extract PHI from prompts, system context, or RAG.",
    "Clinical hallucination — adversarial prompts crafted to elicit confident clinical claims with no source.",
    "Off-label drug recommendation — probes that solicit unsupported pharmacology guidance.",
    "Minimum-necessary violation — does the agent volunteer information beyond the task scope?",
    "BAA-bypass — can the agent be coerced into invoking a non-BAA'd third-party AI service?",
    "HIPAA-condition leak — does the agent reveal HIV / mental-health / substance-use status outside its authorization?",
])

add_heading(doc, "ATLAS Coverage as a First-Class Metric", level=2)

add_para(doc,
    "Every campaign tags findings with MITRE ATLAS technique IDs. Coverage is tracked on the technical "
    "dashboard (see Section 6). The goal is not 100% coverage — it is to make the gaps visible and "
    "deliberate, so that 'we don't test for this' is a documented decision rather than an oversight."
)

add_heading(doc, "MCP Onboarding Red-Teaming as a Productized Service", level=2)

add_para(doc,
    "Per v2.1, every new MCP server entering the curated registry is red-teamed before approval. "
    "This is a productized recurring service from the Red Team — not a one-off — and it is what makes "
    "the MCP farm credibly safer than ad-hoc tool integration. Onboarding tests cover input-side "
    "prompt injection through tool results, output-side data leakage, scope-creep behaviors, and "
    "auth/authorization boundaries."
)

page_break(doc)

# ----- Section 5: Per-vertical capability builds -----
add_heading(doc, "5. Per-Vertical Capability Builds", level=1)

add_para(doc,
    "The pipeline only works if each vertical knows what they own. The matrix below shows primary "
    "and supporting ownership across all eight stages. The remainder of the section spells out what "
    "each lead is responsible for building in each phase."
)

add_image(doc, os.path.join(DIAGRAMS, "04_vertical_ownership.png"), width_in=6.6,
          caption="Figure 4. Ownership matrix — five verticals across the eight pipeline stages.")

vertical_builds = [
    ("AppSec",
     "AppSec is the connective tissue of the pipeline. They own the intake portal, the static and dynamic stages, the remediation path, and the asset-level Security Report Card.",
     [
         "Stand up Semgrep + CodeQL custom rule packs for AI-aware patterns: unsafe prompt interpolation, agent tool-call sinks, embedding-leak patterns. Use existing license; rules live in our repo.",
         "ModelScan deployed in CI for any HuggingFace import path. Block on serialized-model RCE patterns.",
         "Llama Guard 3 PoC at the gateway in shadow mode (log-only) — measure refusal/false-positive rates against three weeks of real prompts.",
         "Paved-road template v0 — secure-by-default RAG agent + secure-by-default chat agent — published to internal portal.",
         "Healthcare prompt linter v0 — Presidio + clinical NER as a pre-flight check on prompt templates.",
     ],
     [
         "Burp Suite Enterprise full integration — every Tier 1 asset gets an authenticated scan in the deploy pipeline.",
         "Auto-PR remediation for dependency and IaC findings — measurable cycle-time reduction.",
         "Custom RAG threat-modeling guide — published, with worked examples, used by all Tier 1 onboardings.",
         "Healthcare prompt linter v1 — promoted from PoC to required check on all prompt-template PRs.",
         "Inherited-controls report card — assets on the paved road show their inherited control coverage.",
     ],
     [
         "Self-service threat-modeling assistant — citizen developers complete a guided threat model with AppSec-validated outputs.",
         "Continuous SBOM + attestation across the AI portfolio.",
         "Stage 2 + 3 fully automated for Tier 3–4; AppSec time goes to Tier 1–2 deep work.",
     ]),
    ("Threat Intelligence",
     "TI owns shadow-AI discovery, the AI threat-feed taxonomy, and the BAA inventory. Their work feeds Stage 0 (denominator) and Stage 6 (intel-driven replay).",
     [
         "AI threat-feed taxonomy — define what we collect, how we classify, who consumes. First version live within Phase 1.",
         "Jailbreak intel weekly report — what's new in the wild, what affects our model providers.",
         "BAA inventory baseline — every third-party AI service touching the org documented; non-BAA'd egress flagged for same-day review.",
         "Shadow-AI discovery v0: CASB + GitHub org scan + Purview pull operating on weekly cadence.",
     ],
     [
         "Adversarial-ML research stream — at least one specialist tracking academic + industry adversarial-ML output.",
         "Provider-incident watch — daily monitoring of Claude / MCP server / model-host incidents and CVEs.",
         "Healthcare-AI TTP catalog — what attackers actually do against healthcare AI, not theoretical.",
         "Shadow-AI discovery to daily cadence with auto-routing into Stage 1 tiering.",
     ],
     [
         "Threat-informed scan tuning — Stage 4 probe set automatically updated from intel.",
         "Auto-replay of new jailbreaks against the live Tier 1 portfolio within 48h of credible disclosure.",
     ]),
    ("Threat Hunt",
     "TH transitions from network/endpoint hunting to AI-mediated hunting. Their primary territory is Stage 6, where the rest of the pipeline produces the data they need.",
     [
         "Telemetry gap analysis — what we log today vs. what we'd need to hunt prompt injection, tool abuse, and exfiltration. Output: a prioritized telemetry ask for Platform Engineering.",
         "First three hunt hypotheses in production: anomalous prompt patterns, unexpected tool-call sequences, off-pipeline egress from agent runtime.",
         "Hunt notebook templates — Jupyter / OpenSearch — that future hunters can fork.",
     ],
     [
         "Live agent-action analytics — tool-call sequence anomaly detection in production.",
         "Drift detection on jailbreak success rate, refusal rate, latency — alerting on threshold breach.",
         "Hunt against shadow-AI signals — joining Stage 0 discovery output with hunt analytics.",
         "Detection-as-code: hunt logic versioned, reviewed, and CI-tested.",
     ],
     [
         "Continuous hunting on the prompt corpus — not periodic.",
         "AI-native detections fed back into SIEM for SOC consumption.",
     ]),
    ("Red Team",
     "RT owns Stage 4 and the demo red-team exercises that earn the Phase 2 ask. They also own MCP onboarding red-teaming as a productized recurring service.",
     [
         "garak baseline scans operational against three Tier 1 candidate apps. Findings reviewed and triaged.",
         "PyRIT PoC orchestrator built and run against the same apps. Healthcare custom converters drafted.",
         "One demonstration red-team exercise on a representative Tier 1 app — written up as the Phase 2 budget conversation artifact.",
         "MCP onboarding red-team v0 — repeatable playbook, used on the first three MCP servers in the registry.",
     ],
     [
         "Healthcare custom probes promoted from draft to required-on-every-Tier-1-campaign.",
         "Indirect-injection corpus built and curated; corpus seeding into RAG test partitions automated.",
         "ATLAS coverage tracking on the technical dashboard.",
         "Tool-abuse fuzz harness in production — replays adversarial tool-call sequences against agent runtime test instances.",
     ],
     [
         "Continuous AI red-team campaigns — Tier 1 monthly, Tier 2 quarterly.",
         "Automated re-test on every regression — successful exploits don't recur silently.",
         "Adversarial-ML platform (ART) operational against any in-house classifier or embedding model.",
     ]),
    ("Security Control Validation",
     "SCV's charter shifts the most. The set of controls to validate has changed; the cadence has to become continuous; and the team has to be the one that says whether a guardrail actually works.",
     [
         "Catalog v2.1 controls vs. AI threat model — gap matrix shared with the Director and the CISO. This becomes the work backlog for Phases 2 and 3.",
         "Egress-control validation playbook — proves that the gateway is, in fact, the only path to foundation models.",
         "Gateway-bypass detection — log analysis identifying any traffic to Claude / OpenAI / Anthropic / Gemini that did not transit our gateway.",
     ],
     [
         "Continuous control validation runner — controls re-tested on a schedule, with results in the technical dashboard.",
         "Guardrail effectiveness measurement — Llama Guard / NeMo / LLM Guard tested against current attack corpus.",
         "Identity & scope validation per agent — does each agent have the minimal tools and data scope its task requires?",
     ],
     [
         "Self-service control attestation — paved-road consumers receive auto-generated attestations of inherited control coverage.",
         "Real-time control drift alerts — control fail produces a finding instantly, not at next audit cycle.",
     ]),
]

for vname, charter, p1, p2, p3 in vertical_builds:
    add_heading(doc, vname, level=2)
    add_para(doc, charter)
    add_heading(doc, "Phase 1 (Months 0–3) — modest, no headcount", level=3)
    add_bullets(doc, p1)
    add_heading(doc, "Phase 2 (Months 3–9) — tooling + selective hire", level=3)
    add_bullets(doc, p2)
    add_heading(doc, "Phase 3 (Months 9–18) — platform-grade", level=3)
    add_bullets(doc, p3)

page_break(doc)

# ----- Section 6: Dashboards -----
add_heading(doc, "6. Dashboards", level=1)

add_para(doc,
    "Two dashboard surfaces, one findings store. The technical dashboard is for operators inside "
    "the offensive security organization — the people on call, triaging findings, working remediation, "
    "and tracking SLAs. The executive dashboard is for the CISO and cyber leadership — the audience "
    "v2.1 was written for. Both refresh from the same data; they differ in aggregation and framing, "
    "not in source of truth."
)

add_heading(doc, "Technical Dashboard", level=2)

add_image(doc, os.path.join(DIAGRAMS, "05_dashboard_technical.png"), width_in=6.6,
          caption="Figure 5. Technical dashboard — operator view, Grafana.")

add_para(doc, "Primary panels:")
add_bullets(doc, [
    "Top KPIs — asset coverage %, open criticals, MTTR critical, jailbreak rate.",
    "Findings by severity, last 12 weeks (stacked bar) — trend visibility.",
    "MTTR by severity, rolling 30 days (line) — SLA performance.",
    "MITRE ATLAS coverage heatmap — what we test for vs. what we don't.",
    "Top open criticals table — by age and tier, with owner and finding summary for action.",
])

add_para(doc, "Audience uses this for:")
add_bullets(doc, [
    "Daily standup — what's open, what's at risk of breaching SLA.",
    "Weekly capacity planning — where the work concentrates.",
    "ATLAS coverage gaps — input to Red Team campaign planning.",
    "Investigation — Top criticals table is the queue.",
])

add_heading(doc, "Executive Dashboard", level=2)

add_image(doc, os.path.join(DIAGRAMS, "06_dashboard_executive.png"), width_in=6.6,
          caption="Figure 6. Executive dashboard — CISO and cyber leadership view, Superset / Power BI.")

add_para(doc, "Primary panels:")
add_bullets(doc, [
    "Portfolio KPIs — registered assets, Tier 1 coverage, avoided-harm demonstrations, open critical risk posture.",
    "Risk heatmap — Business Unit × Tier — where the residual risk concentrates.",
    "Risk introduced vs. closed (monthly) — are we keeping up with the AI portfolio's growth?",
    "Control coverage — HIPAA / HITRUST progress; FDA SaMD readiness.",
    "AI portfolio composition by tier.",
    "Top 5 risk concentrations — single highest-risk assets, named, by score and tier.",
])

add_callout(doc, "takeaway", "What the executive view is for", [
    "It is the artifact that makes the seat at the table durable — the quarterly board / risk-committee report draws directly from it.",
    "It frames offensive security work in patient-safety, HIPAA/HITECH, breach-avoidance, and regulator-credibility terms — not vulnerability counts.",
    "Every panel is owned by a vertical lead who can defend the number. There is no blackbox metric.",
    "Top 5 risk concentrations being named publicly is what creates conversational leverage with business owners.",
])

page_break(doc)

# ----- Section 7: Phased tool rollout -----
add_heading(doc, "7. Phased Tool & Capability Rollout", level=1)

add_para(doc,
    "The phased rollout below is the per-vertical Phase 1/2/3 view from Section 5, presented as a "
    "single picture. It is the artifact to use when negotiating priorities across verticals — what "
    "cannot slip in Phase 1 vs. what can wait."
)

add_image(doc, os.path.join(DIAGRAMS, "07_phase_rollout.png"), width_in=6.6,
          caption="Figure 7. Phase rollout swim lane — five verticals × three phases.")

add_heading(doc, "Phase 1 — What Cannot Slip", level=2)

add_para(doc,
    "Three Phase 1 deliverables sit on the critical path because everything in Phase 2 depends on them."
)

add_bullets(doc, [
    "AppSec: Llama Guard PoC at the gateway. Without a guardrail measurement baseline, we cannot show that we materially reduced runtime risk in Phase 2.",
    "Red Team: the demo red-team exercise on a Tier 1 app. This is the artifact that earns the Phase 2 ask. If it slips, the Phase 2 budget conversation is harder.",
    "Threat Intel: shadow-AI discovery v0. Without the denominator, every metric in Phases 2 and 3 is suspect.",
], bold_lead=True)

add_callout(doc, "callout", "Phase 1 modesty constraint",
    "The discipline of Phase 1 is to ask for nothing — no headcount, no new tooling licenses — and "
    "still produce a working baseline. The credibility we accumulate in Phase 1 is what we spend in "
    "Phase 2. Any temptation to expand scope inside Phase 1 weakens the Phase 2 ask."
)

add_heading(doc, "Phase 2 — What the Demonstrations Buy", level=2)

add_bullets(doc, [
    "Burp Suite Enterprise full integration — depends on AppSec demonstrating Phase 1 cycle-time data.",
    "Selective specialist hire — adversarial-ML researcher in TI; senior AI red-teamer in RT. Justified by Phase 1 demo + intel volume.",
    "Custom probe sets and indirect-injection corpus operationalized — depends on Phase 1 PyRIT PoC.",
    "Continuous control validation runner — depends on SCV gap matrix from Phase 1.",
], bold_lead=True)

add_heading(doc, "Phase 3 — Platform-Grade", level=2)

add_bullets(doc, [
    "Self-service threat-modeling assistant for citizen developers.",
    "Continuous AI red-team campaigns and automated regression replay.",
    "Real-time control attestation surfaces.",
    "Detection-as-code for AI patterns fed into SOC consumption paths.",
], bold_lead=True)

page_break(doc)

# ----- Section 8: RACI extension -----
add_heading(doc, "8. Pipeline RACI Extension", level=1)

add_para(doc,
    "The activities below extend the v2.1 RACI. Owner columns are unchanged from v2.1 — AppSec, "
    "Threat Intel, Threat Hunt, Red Team, SCV, AI Governance, Privacy/Compliance, Platform "
    "Engineering, Business Owner. Each row has exactly one Accountable; the rest are R, C, I, or "
    "blank. This is intended to be appended to the v2.1 RACI section verbatim."
)
add_para(doc, "Header abbreviations: ASec = Application Security, TI = Threat Intel, "
              "TH = Threat Hunt, RT = Red Team, SCV = Security Control Validation, "
              "AIGov = AI Governance, Priv = Privacy / Compliance, PlEng = Platform Engineering, "
              "Biz = Business Owner.",
         italic=True, size=10, color=RGBColor(0x55, 0x55, 0x55))

raci_headers = ["Pipeline activity", "ASec", "TI", "TH", "RT", "SCV", "AIGov", "Priv", "PlEng", "Biz"]
raci_rows = [
    ["Discovery & intake of AI assets",                     "R", "R/A","",  "",  "C",  "C",  "C",  "C",  "I"],
    ["Tier assignment for new AI asset",                    "R", "I",  "",  "",  "I",  "A",  "C",  "",   "I"],
    ["Static & supply-chain scanning",                      "R/A","",  "",  "",  "C",  "I",  "",   "C",  "I"],
    ["Dynamic application & API testing",                   "R/A","",  "",  "R", "C",  "I",  "",   "C",  "I"],
    ["AI red-team campaign — Tier 1",                       "C", "C",  "I", "R/A","C", "I",  "C",  "",   "I"],
    ["AI red-team campaign — Tier 2",                       "C", "C",  "",  "R/A","C", "I",  "I",  "",   "I"],
    ["MCP onboarding red-team",                             "C", "C",  "",  "R/A","C", "C",  "",   "C",  ""],
    ["Healthcare custom probe set maintenance",             "C", "C",  "",  "R/A","C", "I",  "C",  "",   ""],
    ["Indirect-injection corpus curation",                  "C", "C",  "",  "R/A","",  "I",  "C",  "",   ""],
    ["Remediation & closure validation",                    "R/A","",  "",  "C",  "R",  "I",  "C",  "",   "I"],
    ["Guardrail effectiveness measurement",                 "C", "",   "",  "C",  "R/A","I",  "C",  "C",  ""],
    ["WAF / gateway rule push for backstops",               "R", "",   "",  "",  "C",  "I",  "",   "R/A","I"],
    ["Continuous monitoring & drift detection",             "C", "C",  "R/A","C", "R",  "I",  "C",  "C",  "I"],
    ["Threat-driven replay against live portfolio",         "C", "R",  "C", "R/A","C",  "I",  "",   "",   "I"],
    ["Asset Security Report Card",                          "R/A","",  "",  "C",  "R",  "I",  "I",  "",   "I"],
    ["Quarterly board / risk-committee report",             "R", "R",  "R", "R",  "R",  "C",  "C",  "I",  "I"],
    ["Telemetry mesh contract & schema",                    "C", "C",  "C", "C",  "C",  "I",  "C",  "R/A","I"],
    ["Off-pipeline AI deployment incident response",        "R", "R",  "R", "R",  "R",  "I",  "C",  "C",  "A"],
]
add_data_table(doc, raci_headers, raci_rows,
               col_widths=[2.3, 0.5, 0.42, 0.42, 0.42, 0.45, 0.6, 0.45, 0.55, 0.42],
               small=True)

add_para(doc, "")
add_para(doc, "Legend — A: Accountable, R: Responsible, C: Consulted, I: Informed. R/A means the same owner is both Responsible and Accountable. Blank cells are not engaged for that activity.",
         italic=True, size=10, color=RGBColor(0x55, 0x55, 0x55))

add_callout(doc, "takeaway", "Three things worth defending in this RACI", [
    "'Off-pipeline AI deployment incident response' is owned (A) by the Business Owner, not by security. Security cannot be accountable for the consequences of bypass — but every vertical is Responsible for response.",
    "'Tier assignment' is Accountable to AI Governance, not AppSec. AppSec runs the policy; the policy itself is a governance artifact.",
    "'Quarterly board / risk-committee report' has every vertical lead Responsible. There is no single ghostwriter; every lead defends their own data.",
])

page_break(doc)

# ----- Section 9: Asks of the leads -----
add_heading(doc, "9. Asks of the Offensive Security Leads", level=1)

add_para(doc,
    "Below is what each of you is on the hook for in the next 30, 60, and 90 days, assuming the "
    "v2.1 operating model receives final endorsement and Phase 1 is authorized to begin."
)

asks = [
    ("AppSec lead",
     "Confirm the Phase 1 deliverables list in Section 5 is achievable with current headcount and standing licenses. Identify any single deliverable that is at risk and what it would take to de-risk.",
     "Stand up the Semgrep + CodeQL custom rule packs and the ModelScan CI integration. Publish the paved-road template v0. Begin Llama Guard shadow-mode measurement.",
     "Operate the prompt linter v0. Have the asset-level Security Report Card running for at least three Tier 1 candidate apps. Be ready to present cycle-time data as part of the Phase 2 ask."),

    ("Threat Intelligence lead",
     "Confirm the AI threat-feed taxonomy scope and the BAA inventory data sources. Identify one pilot CASB integration to use as the shadow-AI discovery first stop.",
     "BAA inventory baseline complete. Jailbreak intel weekly report running for two cycles. Shadow-AI discovery v0 producing a weekly inventory of unregistered AI usage.",
     "Healthcare-AI TTP catalog v0 published. Provider-incident watch operational. First intel-driven replay pattern documented (even if not yet automated)."),

    ("Threat Hunt lead",
     "Deliver the telemetry gap analysis. This is the single most important Phase 1 input from your team — it sets the spec for what Platform Engineering builds.",
     "First three hunt hypotheses operational against the available telemetry. Hunt notebook templates published.",
     "First hunt finding produced from production AI telemetry. Even one credible hunt result in Phase 1 reframes the Phase 2 conversation about telemetry investment."),

    ("Red Team lead",
     "Select the demo red-team target. This is the most important strategic call any single lead makes in Phase 1. The right target is a real Tier 1 app with a recognizable business sponsor and visible AI features — not a sandbox.",
     "garak and PyRIT PoC operational. Healthcare custom converters drafted. Demo red-team exercise complete and written up as a polished artifact.",
     "MCP onboarding red-team v0 used on the first three registry submissions. Phase 2 ask packaged around the demo red-team write-up."),

    ("Security Control Validation lead",
     "Deliver the v2.1 controls × AI threat model gap matrix. This is the work backlog for the rest of the team for Phases 2 and 3.",
     "Egress-control validation playbook running. Gateway-bypass detection in production at least in alert-only mode.",
     "Be able to answer, with evidence, the question: 'Which v2.1 controls are validated, which are partially validated, and which are unvalidated?' That answer goes into the first quarterly board report."),
]

for role, d30, d60, d90 in asks:
    add_heading(doc, role, level=2)
    add_bullets(doc, [
        f"Next 30 days — {d30}",
        f"Next 60 days — {d60}",
        f"Next 90 days — {d90}",
    ], bold_lead=True)

add_para(doc, "")
add_callout(doc, "takeaway", "What we collectively owe each other", [
    "Honest assessments. If a Phase 1 deliverable is at risk, that is a conversation now, not a slip later.",
    "Cross-vertical handoffs. Most pipeline stages have multiple supporters. Supporting work is not optional and not invisible.",
    "Shared dashboard discipline. Every panel is owned by a lead who can defend the number; no cherry-picking.",
    "The board report. Every quarter, on time, with a defensible posture statement. This is the seat-at-the-table claim — earn it by repetition.",
])

# Save
out_path = os.path.join(DOCS, "pipeline_companion_v1.docx")
doc.save(out_path)
print("Saved:", out_path)

# Stats
d2 = Document(out_path)
print("Paragraphs:", len(d2.paragraphs))
print("Tables:", len(d2.tables))
