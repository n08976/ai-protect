"""Build the v2.1 one-pager and executive variant.

One-pager: dense single-page summary for distribution to peers, AI governance,
privacy, platform engineering. Mirrors v2.1 exactly in framing.

Executive variant: tighter single page for CISO use upward — board, risk
committee, peer execs. Leads with the shift and the ask.

Visual style mirrors v2.1 / companion: Calibri, navy headings, takeaway and
callout shaded tables, navy-banded data table.
"""

import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS = os.path.join(REPO, "docs")
os.makedirs(DOCS, exist_ok=True)

# ---- palette (matches v2.1) ----
NAVY_RGB = RGBColor(0x1F, 0x3A, 0x5F)
WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)
GREY_RGB = RGBColor(0x55, 0x55, 0x55)
ACCENT_RGB = RGBColor(0xC0, 0x4A, 0x2B)
NAVY = "1F3A5F"
TAKEAWAY = "EAF3FA"
CALLOUT = "FFF4E5"
ALT_ROW = "F2F5F9"


# ---------- low-level helpers ----------

def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_cell_borders(cell, color="BFC8D6", size="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=80, left=140, bottom=80, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        m = OxmlElement(f"w:{tag}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def fixed_layout(t, col_widths_in, tight=True):
    tbl = t._tbl
    tblPr = tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    if tight:
        for old in tblPr.findall(qn("w:tblCellMar")):
            tblPr.remove(old)
        tcm = OxmlElement("w:tblCellMar")
        for tag, val in (("top", "30"), ("bottom", "30"), ("left", "70"), ("right", "70")):
            m = OxmlElement(f"w:{tag}")
            m.set(qn("w:w"), val)
            m.set(qn("w:type"), "dxa")
            tcm.append(m)
        tblPr.append(tcm)
    total_twips = int(sum(col_widths_in) * 1440)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(total_twips))
    tblW.set(qn("w:type"), "dxa")
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblPr.append(tblW)
    for existing in tbl.findall(qn("w:tblGrid")):
        tbl.remove(existing)
    grid = OxmlElement("w:tblGrid")
    for w in col_widths_in:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w * 1440)))
        grid.append(gc)
    tblPr.addnext(grid)
    for row in t.rows:
        for ci, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(int(col_widths_in[ci] * 1440)))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


def configure(doc, top=0.5, bottom=0.5, left=0.6, right=0.6):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(2)
    for lvl in range(1, 4):
        try:
            h = doc.styles[f"Heading {lvl}"]
            h.font.color.rgb = NAVY_RGB
            h.font.name = "Calibri"
        except KeyError:
            pass
    for sec in doc.sections:
        sec.top_margin = Inches(top)
        sec.bottom_margin = Inches(bottom)
        sec.left_margin = Inches(left)
        sec.right_margin = Inches(right)


def title_block(doc, eyebrow, title, subtitle, *, title_size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(eyebrow.upper())
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = ACCENT_RGB

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(title_size)
    r.font.color.rgb = NAVY_RGB

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(subtitle)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = GREY_RGB

    # accent rule
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C04A2B")
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_head(doc, text, *, size=11.5, space_before=4, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = NAVY_RGB
    return p


def body(doc, text, *, size=10.5, space_after=2, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    if italic:
        r.italic = True
    return p


def bullet(doc, text, *, size=10, lead_bold=True, space_after=1, indent=0.18):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(space_after)
    r0 = p.add_run("•  ")
    r0.font.size = Pt(size)
    r0.font.color.rgb = ACCENT_RGB
    if lead_bold and " — " in text:
        head, tail = text.split(" — ", 1)
        r1 = p.add_run(head)
        r1.bold = True
        r1.font.size = Pt(size)
        r2 = p.add_run(" — " + tail)
        r2.font.size = Pt(size)
    else:
        r = p.add_run(text)
        r.font.size = Pt(size)
    return p


def callout_box(doc, kind, title, lines, *, width_in=7.3, body_size=9, title_size=10, no_spacer=False):
    color = TAKEAWAY if kind == "takeaway" else CALLOUT
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    fixed_layout(t, [width_in], tight=True)
    cell = t.rows[0].cells[0]
    set_cell_shading(cell, color)
    set_cell_borders(cell, color="BFC8D6", size="4")
    set_cell_margins(cell, top=40, left=120, bottom=40, right=120)
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(1)
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(title_size)
    r1.font.color.rgb = NAVY_RGB
    if isinstance(lines, str):
        lines = [lines]
    for ln in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        if len(lines) > 1:
            r = p.add_run("•  " + ln)
        else:
            r = p.add_run(ln)
        r.font.size = Pt(body_size)
    if not no_spacer:
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.space_before = Pt(0)
    return t


def data_table(doc, headers, rows, col_widths, *, head_size=9.5, body_size=9):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    fixed_layout(t, col_widths, tight=True)
    for ci, h in enumerate(headers):
        c = t.rows[0].cells[ci]
        set_cell_shading(c, NAVY)
        set_cell_borders(c, color="FFFFFF", size="4")
        set_cell_margins(c, top=40, bottom=40, left=80, right=80)
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(head_size)
        r.font.color.rgb = WHITE_RGB
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            shade = ALT_ROW if ri % 2 == 0 else "FFFFFF"
            set_cell_shading(c, shade)
            set_cell_borders(c, color="DDE3EB", size="4")
            set_cell_margins(c, top=40, bottom=40, left=80, right=80)
            c.text = ""
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            r.font.size = Pt(body_size)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)
    return t


def two_column(doc, col_widths_in, render_left, render_right):
    """A two-column layout via a single 1x2 table. Each render function takes a cell."""
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    fixed_layout(t, col_widths_in, tight=True)
    for c in t.rows[0].cells:
        # invisible borders
        tcPr = c._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "nil")
            tcBorders.append(b)
        tcPr.append(tcBorders)
        set_cell_margins(c, top=20, left=40, bottom=20, right=40)
        c.text = ""
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    render_left(t.rows[0].cells[0])
    render_right(t.rows[0].cells[1])
    return t


def cell_head(cell, text, *, size=11, color=NAVY_RGB, space_before=0, space_after=2):
    p = cell.add_paragraph() if cell.paragraphs[0].runs else cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def cell_body(cell, text, *, size=10, italic=False, space_after=2):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    if italic:
        r.italic = True
    return p


def cell_bullet(cell, text, *, size=9.5, lead_bold=True, space_after=1, indent=0.12):
    p = cell.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(space_after)
    r0 = p.add_run("•  ")
    r0.font.size = Pt(size)
    r0.font.color.rgb = ACCENT_RGB
    if lead_bold and " — " in text:
        head, tail = text.split(" — ", 1)
        r1 = p.add_run(head)
        r1.bold = True
        r1.font.size = Pt(size)
        r2 = p.add_run(" — " + tail)
        r2.font.size = Pt(size)
    else:
        r = p.add_run(text)
        r.font.size = Pt(size)


def footer_strip(doc, left_text, right_text, *, width_left=3.65, width_right=3.65):
    # Use a single paragraph with a tab stop for right-aligned text plus a top border.
    # Lighter than a table — avoids the table's vertical overhead that pushes to next page.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    # top border
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")
    top.set(qn("w:space"), "2")
    top.set(qn("w:color"), "C04A2B")
    pBdr.append(top)
    pPr.append(pBdr)
    # tab stop at right edge
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int((width_left + width_right) * 1440)))
    tabs.append(tab)
    pPr.append(tabs)
    r = p.add_run(left_text + "\t" + right_text)
    r.font.size = Pt(8.5)
    r.italic = True
    r.font.color.rgb = GREY_RGB


# ============================================================
# ONE-PAGER — full distribution
# ============================================================

def build_onepager(out_path):
    doc = Document()
    configure(doc, top=0.35, bottom=0.3, left=0.55, right=0.55)

    title_block(
        doc,
        eyebrow="Companion to v2.1 · Internal — for CISO and cyber executive review",
        title="Offensive Security Operating Model for the AI Transformation",
        subtitle="One-page summary — the shift, the model, the infrastructure, the ask.",
    )

    # The shift — a takeaway-style box up top
    callout_box(
        doc, "takeaway", "The shift",
        [
            "From gatekeeper to embedded advisor and continuous validator — the empirical truth function for AI risk.",
            "Pre-deployment review cannot scale to hundreds of citizen-developer apps; the response is to redesign the model, not review harder.",
        ],
        width_in=7.3, body_size=9, title_size=10,
    )

    # Two columns: left = operating principles + tiering, right = infra control plan
    def left(cell):
        cell_head(cell, "Three operating principles", size=10.5)
        cell_bullet(cell, "Paved road over gates — sanctioned gateway, vetted MCP registry, secure-by-default templates. Adoption is the enforcement mechanism.", size=8.5)
        cell_bullet(cell, "Risk-tiered engagement — human depth on Tier 1–2; Tier 3–4 ride automated assurance with sampled red team validation.", size=8.5)
        cell_bullet(cell, "Continuous validation — replace point-in-time pre-deploy review with post-deploy validation, hunting, and control testing.", size=8.5)

        cell_head(cell, "Risk tiering — how effort scales", size=10.5, space_before=4)
        cell_bullet(cell, "Tier 1 — PHI / clinical / external-facing. Embedded AppSec partner, manual red team, continuous control validation.", size=8.5)
        cell_bullet(cell, "Tier 2 — Sensitive internal action / write-back to systems of record. Embedded review, manual red team for material changes.", size=8.5)
        cell_bullet(cell, "Tier 3 — Internal advisory with broad reach. Async checklist review, automated red team, continuous monitoring.", size=8.5)
        cell_bullet(cell, "Tier 4 — Low-impact assistive. Paved-road template, automated scanning, baseline logging, no human review unless flags fire.", size=8.5)
        cell_body(cell, "Dimensions: data sensitivity · decision impact · integration footprint · user population. Re-tier on material change, incident, regulatory shift, or annual recert.", size=8, italic=True, space_after=0)

    def right(cell):
        cell_head(cell, "Sanctioned AI infrastructure — six layers", size=10.5)
        cell_bullet(cell, "AI gateway — only sanctioned path to Claude and approved models; identity, data-class routing, prompt DLP/PHI redaction, output filtering, quotas, full logging.", size=8.5)
        cell_bullet(cell, "MCP server farm — curated registry, no BYO; tier inherited from MCP to calling agent; scoped short-lived tokens; tool-call logging.", size=8.5)
        cell_bullet(cell, "Agent runtime — managed workloads with workload identity, scoped tool allow-list, validated eval suite, sandbox, kill-switch, recert.", size=8.5)
        cell_bullet(cell, "AI-aware SDLC — five stages with automated checks at every tier and gates calibrated to tier (intake → design → build → pre-prod → production).", size=8.5)
        cell_bullet(cell, "Network provisioning — per-tier subnets, mTLS east-west, TLS-inspecting egress, default-deny allow-list, end-user shadow-AI block.", size=8.5)
        cell_bullet(cell, "Unified AI telemetry — prompts, completions, tool calls, retrievals, agent decisions, policy events with identity context — to SIEM, hunt, control validation, owner dashboards.", size=8.5)

    two_column(doc, [3.65, 3.65], left, right)

    # Phased roadmap as a compact 4-column data table
    section_head(doc, "Phased capability roadmap (18 months)", size=10.5, space_before=2, space_after=1)
    data_table(
        doc,
        headers=["Phase 1 — Demonstrate", "Phase 2 — Build the paved road", "Phase 3 — Measure", "Phase 4 — Institutionalize"],
        rows=[[
            "Existing headcount; reprioritize toward AI. Demonstration red team reframes every later budget conversation.",
            "Gateway, MCP farm, agent runtime, AI-aware SDLC. Selective specialist hiring. Paved-road UX as first-class concern.",
            "Continuous control validation for AI; board / risk-committee metrics; sharpened ask informed by data.",
            "Quarterly state-of-AI-security cadence; durable RACI; the seat at the table outlives organizational change.",
        ]],
        col_widths=[1.825, 1.825, 1.825, 1.825],
        head_size=8.5, body_size=8,
    )

    # The ask — a callout
    callout_box(
        doc, "callout", "What we are asking from the CISO and cyber executive team",
        [
            "Operating model approval — gatekeeper → embedded advisor + continuous validator; risk-tiered engagement as the standard.",
            "Risk-tiering adoption — four tiers, socialized with AI governance, privacy, compliance.",
            "AI infrastructure endorsement — gateway, MCP farm, agent runtime, AI-aware SDLC, network provisioning.",
            "RACI ratification — offensive security owns the discipline; platform engineering operates the controls.",
            "Phase 1 authorization — within existing headcount and operating budget; training and tooling PoC funding.",
            "Executive air cover — sponsorship in AI transformation forums, AI governance, and BU leader conversations.",
            "Reporting cadence — quarterly state-of-AI-security report to cyber exec team and risk committee.",
        ],
        width_in=7.3, body_size=8.5, title_size=10, no_spacer=True,
    )

    footer_strip(
        doc,
        left_text="Prepared by the Office of the Director, Offensive Security",
        right_text="Companion summary to v2.1 · 2026-05-08",
    )

    doc.save(out_path)


# ============================================================
# EXECUTIVE VARIANT — for CISO use upward (board / risk committee)
# ============================================================

def build_exec(out_path):
    doc = Document()
    configure(doc, top=0.45, bottom=0.4, left=0.65, right=0.65)

    title_block(
        doc,
        eyebrow="For board, risk committee, and cyber executive leadership",
        title="Offensive Security in the AI Transformation",
        subtitle="Executive brief — why the operating model must shift now, and what we are asking the executive team to endorse.",
    )

    # The thesis as a hero takeaway
    callout_box(
        doc, "takeaway", "The thesis",
        [
            "Within twelve months, hundreds of AI-enabled and self-developed applications will exist across the enterprise — many handling PHI, many deploying outside the formal SDLC.",
            "If offensive security keeps operating as a pre-deployment review function, it will be routed around. If it becomes the empirical truth function for AI risk — the team that proves what does and does not work — it becomes a strategic partner to the transformation.",
            "In healthcare, an AI-mediated PHI exposure or a manipulated clinical-adjacent system carries breach-notification, HIPAA, FDA-adjacent, and patient-safety consequences. The window to shape the model is now.",
        ],
        width_in=7.1, body_size=9, title_size=10,
    )

    # Two columns: what we will do / what we need from leadership
    def left(cell):
        cell_head(cell, "What we will do", size=11)
        cell_bullet(cell, "Build the paved road — sanctioned LLM gateway, curated MCP registry, managed agent runtime, secure-by-default templates. The secure path becomes the easy path; adoption is the enforcement mechanism.", size=9)
        cell_bullet(cell, "Risk-tier every AI initiative — four tiers anchored on PHI, decision impact, integration footprint, and user population. Human depth on Tier 1–2; automated assurance on Tier 3–4.", size=9)
        cell_bullet(cell, "Validate continuously, not once — replace pre-deploy review with continuous post-deploy validation, threat hunting, and control testing across the five offensive security functions.", size=9)
        cell_bullet(cell, "Demonstrate before we ask for more — the Phase 1 demonstration red team exercise reframes every later budget conversation in terms of avoided harm, not abstract risk.", size=9)
        cell_bullet(cell, "Report quarterly — state-of-AI-security to the cyber exec team and risk committee, drawn from unified AI telemetry across gateway, MCP farm, and agent runtime.", size=9)

    def right(cell):
        cell_head(cell, "What we need from the executive team", size=11, color=ACCENT_RGB)
        cell_bullet(cell, "Endorse the operating-model shift — gatekeeper → embedded advisor and continuous validator; risk-tiered engagement as the enterprise standard.", size=9)
        cell_bullet(cell, "Sponsor the sanctioned AI infrastructure — gateway, MCP farm, agent runtime, AI-aware SDLC, network provisioning — in conversations with platform engineering and AI governance.", size=9)
        cell_bullet(cell, "Ratify the RACI — offensive security owns the discipline; platform engineering operates the controls; business owners own tiering inputs and risk acceptance.", size=9)
        cell_bullet(cell, "Authorize Phase 1 — within existing headcount and operating budget, with funded training and tooling proof-of-concepts.", size=9)
        cell_bullet(cell, "Provide air cover — active sponsorship in AI transformation forums, AI governance, and BU leader conversations, positioning offensive security as the empirical truth function for AI risk.", size=9)
        cell_bullet(cell, "Commit the cadence — quarterly state-of-AI-security report to cyber exec team and risk committee, beginning end of Phase 1.", size=9)

    two_column(doc, [3.5, 3.5], left, right)

    # Outcomes — what success looks like
    section_head(doc, "What success looks like — twelve months", size=10.5, space_before=2, space_after=1)
    data_table(
        doc,
        headers=["Velocity preserved", "Risk made visible", "Footing established"],
        rows=[[
            "Citizen developers ship AI tools through a sanctioned path faster and cleaner than shadow alternatives. Hard pre-deploy gates apply only to Tier 1.",
            "Every prompt, tool call, and agent decision flows through unified telemetry. PHI egress, prompt-injection, and agent drift are detectable, attributable, reportable.",
            "Offensive security is a named participant in AI governance, AI infrastructure decisions, and quarterly board reporting — not a downstream reviewer.",
        ]],
        col_widths=[2.4, 2.4, 2.4],
        head_size=9, body_size=8.5,
    )

    # Closing rationale
    callout_box(
        doc, "callout", "Why this proposal, and why now",
        [
            "Phase 1 is executable inside existing headcount — endorsement, not funding, is the unblocking decision.",
            "The Phase 1 demonstration red team is the single highest-leverage move; it reframes every later investment in terms of avoided harm.",
            "The longer AI governance is set without offensive security at the table, the more expensive it is to insert later.",
        ],
        width_in=7.1, body_size=9, title_size=10,
    )

    footer_strip(
        doc,
        left_text="Prepared by the Office of the Director, Offensive Security",
        right_text="Executive brief · companion to v2.1 · 2026-05-08",
        width_left=3.55, width_right=3.55,
    )

    doc.save(out_path)


if __name__ == "__main__":
    one_path = os.path.join(DOCS, "one_pager_v1.docx")
    exec_path = os.path.join(DOCS, "exec_brief_v1.docx")
    build_onepager(one_path)
    build_exec(exec_path)
    print("WROTE:", one_path)
    print("WROTE:", exec_path)
